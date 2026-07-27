"""按扩展名读取日记原文（md/txt/docx）。"""

from __future__ import annotations

from pathlib import Path


def read_file_text(abs_path: str) -> str:
    """读文件正文；.docx 用 python-docx，其它按 UTF-8 文本。"""
    path = Path(abs_path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx(path)
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "读取 .docx 需要 python-docx：pip install python-docx"
        ) from exc
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # 表格单元格一并抽出
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts).strip()
